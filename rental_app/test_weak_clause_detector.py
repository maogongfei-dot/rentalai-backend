# Module3 Phase5-2：弱条款检测 — 最小测试
# 基于 sample.pdf / sample.docx 的 clause_blocks 调用 detect_weak_clauses，打印 weak_clause_count 及每个 weak clause 的 clause_type / page / why_weak。

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from document_reader import read_document
from clause_locator import build_clause_blocks
from weak_clause_detector import detect_weak_clauses, find_blocks_by_keywords, is_deposit_clause_weak

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SAMPLE_DOCX = os.path.join(_SCRIPT_DIR, "sample.docx")
SAMPLE_PDF = os.path.join(_SCRIPT_DIR, "sample.pdf")


def _ensure_samples():
    if not os.path.isfile(SAMPLE_DOCX):
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("Tenancy Agreement – Sample.")
            doc.add_paragraph("The landlord may increase the rent by giving one month notice.")
            doc.add_paragraph("The tenant must pay the deposit to the scheme within 30 days.")
            doc.save(SAMPLE_DOCX)
        except ImportError:
            pass
    if not os.path.isfile(SAMPLE_PDF):
        try:
            from pypdf import PdfWriter
            w = PdfWriter()
            w.add_blank_page(width=595, height=842)
            with open(SAMPLE_PDF, "wb") as f:
                w.write(f)
        except ImportError:
            pass


def _print_weak(label: str, weak: list) -> None:
    print(f"--- {label} ---")
    print("  weak_clause_count:", len(weak))
    for w in weak:
        if isinstance(w, dict):
            print("    clause_type:", w.get("clause_type"), "| page:", w.get("page"), "| why_weak:", w.get("why_weak"))
    print()


def test_weak_clauses_docx():
    """sample.docx 的 clause_blocks 经 detect_weak_clauses 后打印 count 与每个 clause_type / page / why_weak。"""
    _ensure_samples()
    doc = read_document(SAMPLE_DOCX)
    if doc.get("error"):
        print("  skip docx:", doc.get("error"))
        return
    blocks = build_clause_blocks(doc)
    weak = detect_weak_clauses(blocks)
    assert isinstance(weak, list)
    for w in weak:
        assert w.get("clause_type") and w.get("status") == "weak" and w.get("why_weak") is not None
    _print_weak("sample.docx weak_clauses", weak)


def test_weak_clauses_pdf():
    """sample.pdf 的 clause_blocks 经 detect_weak_clauses 后打印 count 与每个 clause_type / page / why_weak。"""
    _ensure_samples()
    doc = read_document(SAMPLE_PDF)
    if doc.get("error"):
        print("  skip pdf:", doc.get("error"))
        return
    blocks = build_clause_blocks(doc)
    weak = detect_weak_clauses(blocks)
    assert isinstance(weak, list)
    _print_weak("sample.pdf weak_clauses", weak)


def test_find_blocks_and_deposit_weak():
    """find_blocks_by_keywords 与 is_deposit_clause_weak 行为正确。"""
    blocks = [
        {"text": "The tenant must pay the deposit to the scheme.", "section_title": "", "page": 1},
        {"text": "Rent is due monthly.", "section_title": "", "page": 2},
    ]
    found = find_blocks_by_keywords(blocks, ["deposit", "scheme"])
    assert len(found) >= 1
    assert is_deposit_clause_weak({"text": "Deposit required.", "section_title": ""}) is True
    assert is_deposit_clause_weak({"text": "The deposit must be protected in a scheme.", "section_title": ""}) is False
    print("find_blocks_by_keywords / is_deposit_clause_weak 测试通过。")


def test_result_integration():
    """build_contract_risk_result_from_document 返回结构含 weak_clauses、weak_clause_count。"""
    from module3_risk_result import build_contract_risk_result_from_document
    _ensure_samples()
    doc = read_document(SAMPLE_DOCX)
    if doc.get("error"):
        print("  skip integration (doc error):", doc.get("error"))
        return
    result = build_contract_risk_result_from_document(doc)
    assert "weak_clauses" in result
    assert "weak_clause_count" in result
    assert result["weak_clause_count"] == len(result["weak_clauses"])
    print("build_contract_risk_result_from_document 含 weak_clauses / weak_clause_count，测试通过。")


if __name__ == "__main__":
    print("Phase5-2 weak_clause_detector 最小测试\n")
    test_find_blocks_and_deposit_weak()
    test_weak_clauses_docx()
    test_weak_clauses_pdf()
    test_result_integration()
    print("Phase5-2 测试完成。")
