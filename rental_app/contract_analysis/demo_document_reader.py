"""
本地快速验证：TXT / PDF / DOCX 文本提取（不跑完整合同规则）。
完整分析 + explain 见 ``demo_contract_document_readers`` / ``service.analyze_contract_file_with_explain``。

运行（在 ``rental_app`` 目录下）::

    python -m contract_analysis.demo_document_reader

依赖：``requirements.txt`` 中的 ``pypdf``、``python-docx``。
"""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

# Windows 控制台尽量用 UTF-8，避免中文错误信息乱码
if hasattr(sys.stdout, "reconfigure"):
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

from .contract_document_reader import (
    MINIMAL_CONTRACT_PDF_TEXT_BYTES,
    ContractReadOutcome,
    extract_contract_text_outcome,
    read_contract_from_docx_outcome,
    read_contract_from_pdf_outcome,
    read_contract_from_txt_outcome,
)


def _print_outcome(label: str, o: ContractReadOutcome) -> None:
    err = o.get("error")
    text = (o.get("text") or "").strip()
    if err:
        print(f"  [{label}] error: {err}")
    else:
        preview = text[:120] + ("..." if len(text) > 120 else "")
        # Windows 控制台常见 GBK，避免 £ 等字符触发 UnicodeEncodeError
        print(f"  [{label}] ok, length={len(text)} preview: {ascii(preview)}")


def run_document_reader_smoke() -> None:
    base = Path(__file__).resolve().parent
    rental_app = base.parent

    print("=== contract_document_reader smoke ===\n")

    # 1) TXT — 仓库内样例
    txt_path = base / "samples" / "sample_contract_safe.txt"
    print(f"1) TXT: {txt_path}")
    _print_outcome("txt", read_contract_from_txt_outcome(txt_path))

    # 2) PDF — 内存最小 PDF
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        f.write(MINIMAL_CONTRACT_PDF_TEXT_BYTES)
        pdf_tmp = Path(f.name)
    try:
        print(f"\n2) PDF: temp {pdf_tmp}")
        _print_outcome("pdf", read_contract_from_pdf_outcome(pdf_tmp))
    finally:
        pdf_tmp.unlink(missing_ok=True)

    # 3) DOCX — 临时生成
    try:
        from docx import Document
    except ImportError:
        print("\n3) DOCX: skipped (python-docx not installed)")
        return

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        docx_tmp = Path(f.name)
    try:
        doc = Document()
        doc.add_paragraph("Tenancy Agreement – smoke test.")
        doc.add_paragraph("The landlord may enter with 24 hours notice.")
        doc.save(docx_tmp)
        print(f"\n3) DOCX: temp {docx_tmp}")
        _print_outcome("docx", read_contract_from_docx_outcome(docx_tmp))
    finally:
        docx_tmp.unlink(missing_ok=True)

    # 4) 统一入口
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        f.write(MINIMAL_CONTRACT_PDF_TEXT_BYTES)
        pdf2 = Path(f.name)
    try:
        print(f"\n4) extract_contract_text_outcome (no source_type): {pdf2}")
        _print_outcome("extract", extract_contract_text_outcome(pdf2))
    finally:
        pdf2.unlink(missing_ok=True)

    # 5) 可选：项目根目录 sample.pdf（若存在）
    legacy = rental_app / "sample.pdf"
    if legacy.is_file():
        print(f"\n5) optional legacy PDF: {legacy}")
        _print_outcome("sample.pdf", read_contract_from_pdf_outcome(legacy))

    print("\n=== done ===")


if __name__ == "__main__":
    run_document_reader_smoke()
