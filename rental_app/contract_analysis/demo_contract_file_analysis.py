"""
Phase 3 Part 4：文件路径 → 合同分析（结构化 + explain）演示。

运行（在 ``rental_app`` 目录下）::

    python -m contract_analysis.demo_contract_file_analysis

固定样例（``sample_contract.txt`` / ``.pdf`` / ``.docx``）请见
``demo_contract_document_readers.run_contract_file_demo``；本模块保留向后兼容别名。
"""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

from .contract_document_reader import MINIMAL_CONTRACT_PDF_TEXT_BYTES
from .demo_contract_document_readers import (
    run_contract_file_demo,
    test_contract_document_readers,
)
from .service import analyze_contract_file_with_explain


def _preview(text: str, n: int = 160) -> str:
    t = (text or "").strip().replace("\n", " ")
    if len(t) > n:
        return t[:n] + "..."
    return t


def run_contract_file_analysis_demo() -> None:
    """向后兼容：与 ``run_contract_file_demo`` 相同（固定样例 + 读取 + 分析 + explain）。"""
    run_contract_file_demo()


def validate_contract_file_analysis_demo() -> None:
    """向后兼容：与 ``test_contract_document_readers`` 相同。"""
    test_contract_document_readers()


def run_contract_file_analysis_demo_legacy_tempfiles() -> None:
    """
    旧版演示：临时 PDF + 临时 DOCX + ``sample_contract_safe.txt``（无固定 sample_contract.* 时可用）。
    """
    base = Path(__file__).resolve().parent
    print("=== analyze_contract_file_with_explain demo (legacy tempfiles) ===\n")

    txt_path = base / "samples" / "sample_contract_safe.txt"
    print(f"1) TXT: {txt_path.name}")
    out = analyze_contract_file_with_explain(file_path=txt_path)
    sa = out["structured_analysis"]
    ex = out["explain"]
    print(f"   meta: {sa.get('meta')}")
    print(f"   summary: {_preview(sa.get('summary', '') or '')}")
    print(f"   overall_conclusion: {_preview(ex.get('overall_conclusion', '') or '')}")
    print()

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        f.write(MINIMAL_CONTRACT_PDF_TEXT_BYTES)
        pdf_tmp = Path(f.name)
    try:
        print(f"2) PDF (minimal temp): {pdf_tmp.name}")
        out2 = analyze_contract_file_with_explain(file_path=pdf_tmp)
        sa2 = out2["structured_analysis"]
        print(f"   meta: {sa2.get('meta')}")
        print(f"   summary: {_preview(sa2.get('summary', '') or '')}")
        print(f"   explain: {_preview(out2['explain'].get('overall_conclusion', '') or '')}")
    finally:
        pdf_tmp.unlink(missing_ok=True)
    print()

    try:
        from docx import Document
    except ImportError:
        print("3) DOCX: skipped (python-docx not installed)")
        print("\n=== done ===")
        return

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        docx_tmp = Path(f.name)
    try:
        doc = Document()
        doc.add_paragraph("Assured shorthold tenancy agreement (excerpt).")
        doc.add_paragraph("Rent: £950 per calendar month. Deposit: £1095 held in DPS.")
        doc.add_paragraph("The landlord may not enter without 24 hours notice except emergency.")
        doc.save(docx_tmp)
        print(f"3) DOCX (temp): {docx_tmp.name}")
        out3 = analyze_contract_file_with_explain(file_path=docx_tmp)
        sa3 = out3["structured_analysis"]
        print(f"   meta: {sa3.get('meta')}")
        print(f"   summary: {_preview(sa3.get('summary', '') or '')}")
        print(f"   explain: {_preview(out3['explain'].get('overall_conclusion', '') or '')}")
    finally:
        docx_tmp.unlink(missing_ok=True)

    print("\n=== done ===")


if __name__ == "__main__":
    run_contract_file_analysis_demo()
