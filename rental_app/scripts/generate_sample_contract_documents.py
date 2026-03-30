#!/usr/bin/env python3
"""
生成 ``contract_analysis/samples/`` 下的二进制样例：

- ``sample_contract.pdf`` — 文本层单页（与 ``MINIMAL_CONTRACT_PDF_TEXT_BYTES`` 一致）
- ``sample_contract.docx`` — 若干段落，与 ``sample_contract.txt`` 语义相近

在 ``rental_app`` 目录执行::

    python scripts/generate_sample_contract_documents.py

依赖：``pypdf``（仅复制字节时可省略）、``python-docx``。
"""

from __future__ import annotations

import sys
from pathlib import Path


def _root() -> Path:
    return Path(__file__).resolve().parent.parent


def main() -> None:
    root = _root()
    if str(root) not in sys.path:
        sys.path.insert(0, str(root))

    from contract_analysis.contract_document_reader import MINIMAL_CONTRACT_PDF_TEXT_BYTES

    samples = root / "contract_analysis" / "samples"
    samples.mkdir(parents=True, exist_ok=True)

    pdf_path = samples / "sample_contract.pdf"
    pdf_path.write_bytes(MINIMAL_CONTRACT_PDF_TEXT_BYTES)
    print(f"Wrote {pdf_path} ({len(MINIMAL_CONTRACT_PDF_TEXT_BYTES)} bytes)")

    try:
        from docx import Document
    except ImportError as e:
        print("Skip DOCX: python-docx not installed:", e)
        return

    docx_path = samples / "sample_contract.docx"
    doc = Document()
    doc.add_heading("Assured Shorthold Tenancy — Sample (Part 4)", level=1)
    doc.add_paragraph(
        "This is a minimal demo contract document for RentalAI reader tests. "
        "Rent: £850 per calendar month, payable in advance on the 1st."
    )
    doc.add_paragraph(
        "Deposit: five weeks' rent, to be protected in a government-approved scheme "
        "(DPS / TDS / MyDeposits); prescribed information within 30 days."
    )
    doc.add_paragraph(
        "Notice: landlord not less than two months after fixed term; tenant not less than one month."
    )
    doc.add_paragraph(
        "Repairs: landlord responsible for structure; tenant to report defects promptly."
    )
    doc.add_paragraph(
        "Entry: landlord or agent may enter with at least 24 hours' written notice except emergencies."
    )
    doc.save(docx_path)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    main()
