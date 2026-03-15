# Module3 Phase4 Final：文档风险输出收口 — 最小测试
# 读取 sample.pdf / sample.docx 后，输出 document_summary、clause_blocks 数量、risk_clauses 数量、highlighted_clauses 数量。

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from document_reader import read_document
from module3_risk_result import build_contract_risk_result_from_document

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SAMPLE_DOCX = os.path.join(_SCRIPT_DIR, "sample.docx")
SAMPLE_PDF = os.path.join(_SCRIPT_DIR, "sample.pdf")


def _ensure_samples():
    if not os.path.isfile(SAMPLE_DOCX):
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("Tenancy Agreement – Sample.")
            doc.add_paragraph("The landlord may increase the rent by giving one month notice.")
            doc.add_paragraph("The tenant must pay the deposit to the scheme within 30 days.")
            doc.save(SAMPLE_DOCX)
        except ImportError:
            pass
    if not os.path.isfile(SAMPLE_PDF):
        try:
            from pypdf import PdfWriter
            w = PdfWriter()
            w.add_blank_page(width=595, height=842)
            with open(SAMPLE_PDF, "wb") as f:
                w.write(f)
        except ImportError:
            pass


def test_document_analysis_block_docx():
    """sample.docx 经 build_contract_risk_result_from_document 后，输出 document_summary 与各数量。"""
    _ensure_samples()
    doc = read_document(SAMPLE_DOCX)
    if doc.get("error"):
        print("  skip docx:", doc.get("error"))
        return
    result = build_contract_risk_result_from_document(doc)
    block = result.get("document_analysis_block") or {}
    summary = block.get("document_summary") or {}
    print("--- sample.docx document_analysis_block ---")
    print("  document_summary:", summary)
    print("  clause_blocks count:", len(block.get("clause_blocks") or []))
    print("  risk_clauses count:", len(block.get("risk_clauses") or []))
    print("  highlighted_clauses count:", len(block.get("highlighted_clauses") or []))
    assert summary.get("file_name") == "sample.docx"
    assert summary.get("file_type") == "docx"
    assert summary.get("block_count") == len(block.get("clause_blocks") or [])
    assert summary.get("risk_clause_count") == len(block.get("risk_clauses") or [])
    assert summary.get("highlighted_clause_count") == len(block.get("highlighted_clauses") or [])
    print()


def test_document_analysis_block_pdf():
    """sample.pdf 经 build_contract_risk_result_from_document 后，输出 document_summary 与各数量。"""
    _ensure_samples()
    doc = read_document(SAMPLE_PDF)
    if doc.get("error"):
        print("  skip pdf:", doc.get("error"))
        return
    result = build_contract_risk_result_from_document(doc)
    block = result.get("document_analysis_block") or {}
    summary = block.get("document_summary") or {}
    print("--- sample.pdf document_analysis_block ---")
    print("  document_summary:", summary)
    print("  clause_blocks count:", len(block.get("clause_blocks") or []))
    print("  risk_clauses count:", len(block.get("risk_clauses") or []))
    print("  highlighted_clauses count:", len(block.get("highlighted_clauses") or []))
    assert summary.get("file_name") == "sample.pdf"
    assert summary.get("file_type") == "pdf"
    print()


if __name__ == "__main__":
    print("Phase4 Final document_analysis_builder 最小测试\n")
    test_document_analysis_block_docx()
    test_document_analysis_block_pdf()
    print("Phase4 Final 测试完成。")
