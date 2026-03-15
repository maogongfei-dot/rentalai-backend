# Module3 Phase5-1：缺失条款检测 — 最小测试
# 基于 sample.pdf / sample.docx 的 clause_blocks 调用 detect_missing_clauses，打印 missing_clause_count 及每个 missing clause 的 clause_type / why_missing。

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from document_reader import read_document
from clause_locator import build_clause_blocks
from missing_clause_detector import detect_missing_clauses, has_clause_keywords

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SAMPLE_DOCX = os.path.join(_SCRIPT_DIR, "sample.docx")
SAMPLE_PDF = os.path.join(_SCRIPT_DIR, "sample.pdf")


def _ensure_samples():
    if not os.path.isfile(SAMPLE_DOCX):
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("Tenancy Agreement – Sample.")
            doc.add_paragraph("The landlord may increase the rent by giving one month notice.")
            doc.add_paragraph("The tenant must pay the deposit to the scheme within 30 days.")
            doc.save(SAMPLE_DOCX)
        except ImportError:
            pass
    if not os.path.isfile(SAMPLE_PDF):
        try:
            from pypdf import PdfWriter
            w = PdfWriter()
            w.add_blank_page(width=595, height=842)
            with open(SAMPLE_PDF, "wb") as f:
                w.write(f)
        except ImportError:
            pass


def _print_missing(label: str, missing: list) -> None:
    print(f"--- {label} ---")
    print("  missing_clause_count:", len(missing))
    for m in missing:
        if isinstance(m, dict):
            print("    clause_type:", m.get("clause_type"), "| why_missing:", m.get("why_missing"))
    print()


def test_missing_clauses_docx():
    """sample.docx 的 clause_blocks 经 detect_missing_clauses 后打印 count 与每个 clause_type / why_missing。"""
    _ensure_samples()
    doc = read_document(SAMPLE_DOCX)
    if doc.get("error"):
        print("  skip docx:", doc.get("error"))
        return
    blocks = build_clause_blocks(doc)
    missing = detect_missing_clauses(blocks)
    assert isinstance(missing, list)
    for m in missing:
        assert m.get("clause_type") and m.get("status") == "missing" and m.get("why_missing")
    # sample.docx 含 rent、deposit、notice 等词，故 deposit_clause/rent_clause/notice_clause 应存在，termination/repair 可能缺失
    _print_missing("sample.docx missing_clauses", missing)


def test_missing_clauses_pdf():
    """sample.pdf 的 clause_blocks 经 detect_missing_clauses 后打印 count 与每个 clause_type / why_missing。"""
    _ensure_samples()
    doc = read_document(SAMPLE_PDF)
    if doc.get("error"):
        print("  skip pdf:", doc.get("error"))
        return
    blocks = build_clause_blocks(doc)
    missing = detect_missing_clauses(blocks)
    assert isinstance(missing, list)
    # 空白 PDF 无 blocks，全部 5 类应为缺失
    _print_missing("sample.pdf missing_clauses", missing)


def test_has_clause_keywords():
    """has_clause_keywords 能正确命中 text 中的关键词。"""
    blocks = [{"text": "The tenant must pay the deposit to the scheme.", "section_title": ""}]
    assert has_clause_keywords(blocks, ["deposit", "scheme"]) is True
    assert has_clause_keywords(blocks, ["repair", "maintain"]) is False
    assert has_clause_keywords([], ["deposit"]) is False
    print("has_clause_keywords 测试通过。")


def test_result_integration():
    """build_contract_risk_result_from_document 返回结构含 missing_clauses、missing_clause_count。"""
    from module3_risk_result import build_contract_risk_result_from_document
    _ensure_samples()
    doc = read_document(SAMPLE_DOCX)
    if doc.get("error"):
        print("  skip integration (doc error):", doc.get("error"))
        return
    result = build_contract_risk_result_from_document(doc)
    assert "missing_clauses" in result
    assert "missing_clause_count" in result
    assert result["missing_clause_count"] == len(result["missing_clauses"])
    print("build_contract_risk_result_from_document 含 missing_clauses / missing_clause_count，测试通过。")


if __name__ == "__main__":
    print("Phase5-1 missing_clause_detector 最小测试\n")
    test_has_clause_keywords()
    test_missing_clauses_docx()
    test_missing_clauses_pdf()
    test_result_integration()
    print("Phase5-1 测试完成。")
