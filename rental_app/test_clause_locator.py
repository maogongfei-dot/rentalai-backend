# Module3 Phase4-2：条款切分与定位 — 最小测试
# 读取 sample.pdf / sample.docx 后，用 build_clause_blocks 切分，打印 block 数量及前 2~3 个 block 的 page / section_title / text 前100字符。

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from document_reader import read_document
from clause_locator import build_clause_blocks, segment_pdf_blocks, segment_docx_blocks

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SAMPLE_DOCX = os.path.join(_SCRIPT_DIR, "sample.docx")
SAMPLE_PDF = os.path.join(_SCRIPT_DIR, "sample.pdf")
TEXT_PREVIEW_LEN = 100


def _ensure_samples():
    """若样本不存在则创建（与 test_document_reader 一致）。"""
    if not os.path.isfile(SAMPLE_DOCX):
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("Tenancy Agreement – Sample.")
            doc.add_paragraph("The landlord may increase the rent by giving one month notice.")
            doc.add_paragraph("The tenant must pay the deposit to the scheme within 30 days.")
            doc.save(SAMPLE_DOCX)
        except ImportError:
            pass
    if not os.path.isfile(SAMPLE_PDF):
        try:
            from pypdf import PdfWriter
            w = PdfWriter()
            w.add_blank_page(width=595, height=842)
            with open(SAMPLE_PDF, "wb") as f:
                w.write(f)
        except ImportError:
            pass


def _print_blocks(label: str, blocks: list, max_show: int = 3) -> None:
    """打印 block 数量及前 max_show 个 block 的 page / section_title / text 前100字符。"""
    print(f"--- {label} ---")
    print("  block count:", len(blocks))
    for i, b in enumerate(blocks[:max_show]):
        if not isinstance(b, dict):
            continue
        page = b.get("page")
        title = (b.get("section_title") or "")
        text = (b.get("text") or "")
        preview = text[:TEXT_PREVIEW_LEN] + ("..." if len(text) > TEXT_PREVIEW_LEN else "")
        print(f"  block {i+1}: page={page!r}, section_title={title!r}, text_preview={preview!r}")
    print()


def test_clause_blocks_docx():
    """读取 sample.docx，切分后打印 block 数量与前 2~3 个 block。"""
    _ensure_samples()
    doc = read_document(SAMPLE_DOCX)
    if doc.get("error"):
        print("  skip docx (error or missing file):", doc.get("error"))
        return
    blocks = build_clause_blocks(doc)
    assert isinstance(blocks, list)
    for b in blocks:
        assert "block_id" in b and "page" in b and "section_title" in b and "text" in b
    _print_blocks("sample.docx clause_blocks", blocks, max_show=3)


def test_clause_blocks_pdf():
    """读取 sample.pdf，切分后打印 block 数量与前 2~3 个 block。"""
    _ensure_samples()
    doc = read_document(SAMPLE_PDF)
    if doc.get("error"):
        print("  skip pdf (error or missing file):", doc.get("error"))
        return
    blocks = build_clause_blocks(doc)
    assert isinstance(blocks, list)
    for b in blocks:
        assert "block_id" in b and "page" in b and "section_title" in b and "text" in b
    _print_blocks("sample.pdf clause_blocks", blocks, max_show=3)


def test_build_clause_blocks_dispatch():
    """build_clause_blocks 对含 error 或未知 file_type 返回空列表。"""
    assert build_clause_blocks({}) == []
    assert build_clause_blocks({"error": "x"}) == []
    assert build_clause_blocks({"file_type": "txt"}) == []


if __name__ == "__main__":
    print("Phase4-2 clause_locator 最小测试\n")
    test_build_clause_blocks_dispatch()
    test_clause_blocks_docx()
    test_clause_blocks_pdf()
    print("Phase4-2 测试完成。")
