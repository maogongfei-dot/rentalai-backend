# Module3 Phase4-4：重点条款摘要 — 最小测试
# 基于 sample.pdf / sample.docx 的 risk_clauses 生成 highlighted_clauses，打印 count 及前 2~3 条的 page / section_title / risk_flags / highlight_reason。

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from document_reader import read_document
from clause_locator import build_clause_blocks
from risk_clause_detector import detect_risk_clauses
from highlighted_clause_builder import build_highlighted_clauses

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SAMPLE_DOCX = os.path.join(_SCRIPT_DIR, "sample.docx")
SAMPLE_PDF = os.path.join(_SCRIPT_DIR, "sample.pdf")


def _ensure_samples():
    if not os.path.isfile(SAMPLE_DOCX):
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("Tenancy Agreement – Sample.")
            doc.add_paragraph("The landlord may increase the rent by giving one month notice.")
            doc.add_paragraph("The tenant must pay the deposit to the scheme within 30 days.")
            doc.save(SAMPLE_DOCX)
        except ImportError:
            pass
    if not os.path.isfile(SAMPLE_PDF):
        try:
            from pypdf import PdfWriter
            w = PdfWriter()
            w.add_blank_page(width=595, height=842)
            with open(SAMPLE_PDF, "wb") as f:
                w.write(f)
        except ImportError:
            pass


def _print_highlighted(label: str, highlighted: list, max_show: int = 3) -> None:
    print(f"--- {label} ---")
    print("  highlighted_clause_count:", len(highlighted))
    for i, h in enumerate(highlighted[:max_show]):
        if not isinstance(h, dict):
            continue
        print(f"  highlighted {i+1}: page={h.get('page')!r}, section_title={h.get('section_title')!r}, risk_flags={h.get('risk_flags')!r}, highlight_reason={h.get('highlight_reason')!r}")
    print()


def test_highlighted_docx():
    """sample.docx 的 risk_clauses 经 build_highlighted_clauses 后打印 count 与前 2~3 条。"""
    _ensure_samples()
    doc = read_document(SAMPLE_DOCX)
    if doc.get("error"):
        print("  skip docx:", doc.get("error"))
        return
    blocks = build_clause_blocks(doc)
    risk_clauses = detect_risk_clauses(blocks)
    highlighted = build_highlighted_clauses(risk_clauses, max_items=3)
    assert isinstance(highlighted, list)
    for h in highlighted:
        assert "block_id" in h and "risk_flags" in h and "highlight_reason" in h
    _print_highlighted("sample.docx highlighted_clauses", highlighted, max_show=3)


def test_highlighted_pdf():
    """sample.pdf 的 risk_clauses 经 build_highlighted_clauses 后打印 count 与前 2~3 条。"""
    _ensure_samples()
    doc = read_document(SAMPLE_PDF)
    if doc.get("error"):
        print("  skip pdf:", doc.get("error"))
        return
    blocks = build_clause_blocks(doc)
    risk_clauses = detect_risk_clauses(blocks)
    highlighted = build_highlighted_clauses(risk_clauses, max_items=3)
    assert isinstance(highlighted, list)
    _print_highlighted("sample.pdf highlighted_clauses", highlighted, max_show=3)


def test_result_integration():
    """build_contract_risk_result_from_document 返回结构含 highlighted_clauses、highlighted_clause_count。"""
    from module3_risk_result import build_contract_risk_result_from_document
    _ensure_samples()
    doc = read_document(SAMPLE_DOCX)
    if doc.get("error"):
        print("  skip integration (doc error):", doc.get("error"))
        return
    result = build_contract_risk_result_from_document(doc)
    assert "highlighted_clauses" in result
    assert "highlighted_clause_count" in result
    assert result["highlighted_clause_count"] == len(result["highlighted_clauses"])
    assert result["highlighted_clause_count"] <= 3
    print("build_contract_risk_result_from_document 含 highlighted_clauses / highlighted_clause_count，测试通过。")


if __name__ == "__main__":
    print("Phase4-4 highlighted_clause_builder 最小测试\n")
    test_highlighted_docx()
    test_highlighted_pdf()
    test_result_integration()
    print("Phase4-4 测试完成。")
