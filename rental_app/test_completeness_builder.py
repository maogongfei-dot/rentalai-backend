# Module3 Phase5 Final：缺失条款与完整性输出收口 — 最小测试
# 读取 sample.pdf / sample.docx 后，输出 missing_clause_count、weak_clause_count、contract_completeness、completeness_summary；并校验 completeness_block 结构。

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from document_reader import read_document
from module3_risk_result import build_contract_risk_result_from_document

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SAMPLE_DOCX = os.path.join(_SCRIPT_DIR, "sample.docx")
SAMPLE_PDF = os.path.join(_SCRIPT_DIR, "sample.pdf")


def _ensure_samples():
    if not os.path.isfile(SAMPLE_DOCX):
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("Tenancy Agreement – Sample.")
            doc.add_paragraph("The landlord may increase the rent by giving one month notice.")
            doc.add_paragraph("The tenant must pay the deposit to the scheme within 30 days.")
            doc.save(SAMPLE_DOCX)
        except ImportError:
            pass
    if not os.path.isfile(SAMPLE_PDF):
        try:
            from pypdf import PdfWriter
            w = PdfWriter()
            w.add_blank_page(width=595, height=842)
            with open(SAMPLE_PDF, "wb") as f:
                w.write(f)
        except ImportError:
            pass


def test_completeness_block_docx():
    """sample.docx 经 build_contract_risk_result_from_document 后，输出 missing/weak count、contract_completeness、completeness_summary；校验 completeness_block。"""
    _ensure_samples()
    doc = read_document(SAMPLE_DOCX)
    if doc.get("error"):
        print("  skip docx:", doc.get("error"))
        return
    result = build_contract_risk_result_from_document(doc)
    block = result.get("completeness_block") or {}
    assert "missing_clauses" in block and "weak_clauses" in block and "contract_completeness" in block
    comp = block.get("contract_completeness") or {}
    print("--- sample.docx Phase5 Final ---")
    print("  missing_clause_count:", result.get("missing_clause_count"))
    print("  weak_clause_count:", result.get("weak_clause_count"))
    print("  contract_completeness:", comp)
    print("  completeness_summary:", result.get("completeness_summary"))
    print()


def test_completeness_block_pdf():
    """sample.pdf 经 build_contract_risk_result_from_document 后，输出 missing/weak count、contract_completeness、completeness_summary；校验 completeness_block。"""
    _ensure_samples()
    doc = read_document(SAMPLE_PDF)
    if doc.get("error"):
        print("  skip pdf:", doc.get("error"))
        return
    result = build_contract_risk_result_from_document(doc)
    block = result.get("completeness_block") or {}
    assert "missing_clauses" in block and "weak_clauses" in block and "contract_completeness" in block
    comp = block.get("contract_completeness") or {}
    print("--- sample.pdf Phase5 Final ---")
    print("  missing_clause_count:", result.get("missing_clause_count"))
    print("  weak_clause_count:", result.get("weak_clause_count"))
    print("  contract_completeness:", comp)
    print("  completeness_summary:", result.get("completeness_summary"))
    print()


def test_result_has_required_keys():
    """result 稳定包含所列 Phase5 及关键字段。"""
    _ensure_samples()
    doc = read_document(SAMPLE_DOCX)
    if doc.get("error"):
        print("  skip (doc error):", doc.get("error"))
        return
    result = build_contract_risk_result_from_document(doc)
    required = [
        "input_type", "scenario", "risk_flags", "severity", "explanation",
        "recommended_actions", "action_details", "ordered_action_details",
        "law_topics", "legal_references", "legal_reasoning", "legal_summary",
        "evidence_required", "recommended_steps", "possible_outcomes",
        "clause_blocks", "risk_clauses", "highlighted_clauses",
        "missing_clauses", "weak_clauses", "contract_completeness",
        "completeness_block",
    ]
    for k in required:
        assert k in result, f"result 应包含 {k!r}"
    assert result.get("completeness_block", {}).get("contract_completeness") is not None
    print("result 所需字段及 completeness_block 存在，测试通过。")


if __name__ == "__main__":
    print("Phase5 Final completeness_builder 最小测试\n")
    test_completeness_block_docx()
    test_completeness_block_pdf()
    test_result_has_required_keys()
    print("Phase5 Final 测试完成。")
