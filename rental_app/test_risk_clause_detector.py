# Module3 Phase4-3：风险条款识别 — 最小测试
# 基于 sample.pdf / sample.docx 的 clause_blocks 调用 detect_risk_clauses，打印 risk_clause_count 及前 2~3 条的 page / section_title / risk_flags / why_flagged。

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from document_reader import read_document
from clause_locator import build_clause_blocks
from risk_clause_detector import detect_risk_clauses

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SAMPLE_DOCX = os.path.join(_SCRIPT_DIR, "sample.docx")
SAMPLE_PDF = os.path.join(_SCRIPT_DIR, "sample.pdf")


def _ensure_samples():
    if not os.path.isfile(SAMPLE_DOCX):
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("Tenancy Agreement – Sample.")
            doc.add_paragraph("The landlord may increase the rent by giving one month notice.")
            doc.add_paragraph("The tenant must pay the deposit to the scheme within 30 days.")
            doc.save(SAMPLE_DOCX)
        except ImportError:
            pass
    if not os.path.isfile(SAMPLE_PDF):
        try:
            from pypdf import PdfWriter
            w = PdfWriter()
            w.add_blank_page(width=595, height=842)
            with open(SAMPLE_PDF, "wb") as f:
                w.write(f)
        except ImportError:
            pass


def _print_risk_clauses(label: str, risk_clauses: list, max_show: int = 3) -> None:
    print(f"--- {label} ---")
    print("  risk_clause_count:", len(risk_clauses))
    for i, rc in enumerate(risk_clauses[:max_show]):
        if not isinstance(rc, dict):
            continue
        print(f"  risk clause {i+1}: page={rc.get('page')!r}, section_title={rc.get('section_title')!r}, risk_flags={rc.get('risk_flags')!r}, why_flagged={rc.get('why_flagged')!r}")
    print()


def test_risk_clauses_docx():
    """sample.docx 的 clause_blocks 经 detect_risk_clauses 后打印 count 与前 2~3 条。"""
    _ensure_samples()
    doc = read_document(SAMPLE_DOCX)
    if doc.get("error"):
        print("  skip docx:", doc.get("error"))
        return
    blocks = build_clause_blocks(doc)
    risk_clauses = detect_risk_clauses(blocks)
    assert isinstance(risk_clauses, list)
    for rc in risk_clauses:
        assert "block_id" in rc and "risk_flags" in rc and "risk_level" in rc and "why_flagged" in rc
    _print_risk_clauses("sample.docx risk_clauses", risk_clauses, max_show=3)


def test_risk_clauses_pdf():
    """sample.pdf 的 clause_blocks 经 detect_risk_clauses 后打印 count 与前 2~3 条。"""
    _ensure_samples()
    doc = read_document(SAMPLE_PDF)
    if doc.get("error"):
        print("  skip pdf:", doc.get("error"))
        return
    blocks = build_clause_blocks(doc)
    risk_clauses = detect_risk_clauses(blocks)
    assert isinstance(risk_clauses, list)
    _print_risk_clauses("sample.pdf risk_clauses", risk_clauses, max_show=3)


def test_result_integration():
    """build_contract_risk_result_from_document 返回结构含 risk_clauses、risk_clause_count。"""
    from module3_risk_result import build_contract_risk_result_from_document
    _ensure_samples()
    doc = read_document(SAMPLE_DOCX)
    if doc.get("error"):
        print("  skip integration (doc error):", doc.get("error"))
        return
    result = build_contract_risk_result_from_document(doc)
    assert "risk_clauses" in result
    assert "risk_clause_count" in result
    assert result["risk_clause_count"] == len(result["risk_clauses"])
    assert "clause_blocks" in result
    print("build_contract_risk_result_from_document 含 risk_clauses / risk_clause_count / clause_blocks，测试通过。")


if __name__ == "__main__":
    print("Phase4-3 risk_clause_detector 最小测试\n")
    test_risk_clauses_docx()
    test_risk_clauses_pdf()
    test_result_integration()
    print("Phase4-3 测试完成。")
