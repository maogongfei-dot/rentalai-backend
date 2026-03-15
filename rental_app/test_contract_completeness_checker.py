# Module3 Phase5-3：合同完整性检查 — 最小测试
# 基于 sample.pdf / sample.docx 的 missing_clauses 和 weak_clauses，打印 completeness_level、missing_clause_count、weak_clause_count、completeness_summary。

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from document_reader import read_document
from module3_risk_result import build_contract_risk_result_from_document
from contract_completeness_checker import build_contract_completeness

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SAMPLE_DOCX = os.path.join(_SCRIPT_DIR, "sample.docx")
SAMPLE_PDF = os.path.join(_SCRIPT_DIR, "sample.pdf")


def _ensure_samples():
    if not os.path.isfile(SAMPLE_DOCX):
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("Tenancy Agreement – Sample.")
            doc.add_paragraph("The landlord may increase the rent by giving one month notice.")
            doc.add_paragraph("The tenant must pay the deposit to the scheme within 30 days.")
            doc.save(SAMPLE_DOCX)
        except ImportError:
            pass
    if not os.path.isfile(SAMPLE_PDF):
        try:
            from pypdf import PdfWriter
            w = PdfWriter()
            w.add_blank_page(width=595, height=842)
            with open(SAMPLE_PDF, "wb") as f:
                w.write(f)
        except ImportError:
            pass


def _print_completeness(label: str, comp: dict) -> None:
    print(f"--- {label} ---")
    print("  completeness_level:", comp.get("completeness_level"))
    print("  missing_clause_count:", comp.get("missing_clause_count"))
    print("  weak_clause_count:", comp.get("weak_clause_count"))
    print("  completeness_summary:", comp.get("completeness_summary"))
    print()


def test_completeness_docx():
    """sample.docx 经 build_contract_risk_result_from_document 后打印 contract_completeness。"""
    _ensure_samples()
    doc = read_document(SAMPLE_DOCX)
    if doc.get("error"):
        print("  skip docx:", doc.get("error"))
        return
    result = build_contract_risk_result_from_document(doc)
    comp = result.get("contract_completeness") or {}
    assert "completeness_level" in comp and "completeness_summary" in comp
    _print_completeness("sample.docx contract_completeness", comp)


def test_completeness_pdf():
    """sample.pdf 经 build_contract_risk_result_from_document 后打印 contract_completeness。"""
    _ensure_samples()
    doc = read_document(SAMPLE_PDF)
    if doc.get("error"):
        print("  skip pdf:", doc.get("error"))
        return
    result = build_contract_risk_result_from_document(doc)
    comp = result.get("contract_completeness") or {}
    _print_completeness("sample.pdf contract_completeness", comp)


def test_build_contract_completeness_rules():
    """build_contract_completeness 等级规则：missing>=3->low, missing1~2或weak>=2->medium, 否则 high。"""
    c = build_contract_completeness([{}, {}, {}], [])
    assert c["completeness_level"] == "low"
    c = build_contract_completeness([{}], [{}])
    assert c["completeness_level"] == "medium"
    c = build_contract_completeness([], [{}, {}])
    assert c["completeness_level"] == "medium"
    c = build_contract_completeness([], [])
    assert c["completeness_level"] == "high"
    c = build_contract_completeness([], [{}])
    assert c["completeness_level"] == "high"
    print("build_contract_completeness 等级规则测试通过。")


if __name__ == "__main__":
    print("Phase5-3 contract_completeness_checker 最小测试\n")
    test_build_contract_completeness_rules()
    test_completeness_docx()
    test_completeness_pdf()
    print("Phase5-3 测试完成。")
