# Module3 Phase4-1：Document 读取入口 — 简单测试
# 读取 sample.pdf / sample.docx，打印 file_name、file_type、full_text 前几百字符、pages_or_sections 数量。
# 若样本文件不存在则先创建最小样本再读取。

import os
import sys

# 确保可导入 rental_app 内模块
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from document_reader import read_document, extract_text_from_pdf, extract_text_from_docx

# 样本文件放在 rental_app 目录下
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SAMPLE_DOCX = os.path.join(_SCRIPT_DIR, "sample.docx")
SAMPLE_PDF = os.path.join(_SCRIPT_DIR, "sample.pdf")
PREVIEW_LEN = 400


def _create_sample_docx(path: str) -> None:
    """创建最小 sample.docx（若干段落），便于测试提取。"""
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph("Tenancy Agreement – Sample.")
        doc.add_paragraph("The landlord may increase the rent by giving one month notice.")
        doc.add_paragraph("The tenant must pay the deposit to the scheme within 30 days.")
        doc.save(path)
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")


def _create_sample_pdf(path: str) -> None:
    """创建最小 sample.pdf（一页空白页），便于测试 PDF 读取结构。"""
    try:
        from pypdf import PdfWriter
        writer = PdfWriter()
        writer.add_blank_page(width=595, height=842)
        with open(path, "wb") as f:
            writer.write(f)
    except ImportError:
        raise RuntimeError("pypdf not installed. Run: pip install pypdf")


def _print_result(label: str, result: dict) -> None:
    """打印 file_name、file_type、full_text 前 N 字符、pages_or_sections 数量。"""
    print(f"--- {label} ---")
    print("  file_name:", result.get("file_name"))
    print("  file_type:", result.get("file_type"))
    full = result.get("full_text") or ""
    print("  full_text (first %d chars):" % PREVIEW_LEN, repr(full[:PREVIEW_LEN]) + ("..." if len(full) > PREVIEW_LEN else ""))
    sections = result.get("pages_or_sections") or []
    print("  pages_or_sections count:", len(sections))
    if result.get("error"):
        print("  error:", result["error"])
    print()


def test_read_sample_docx():
    """读取 sample.docx：不存在则创建后读取，并断言统一结构。"""
    if not os.path.isfile(SAMPLE_DOCX):
        _create_sample_docx(SAMPLE_DOCX)
    result = read_document(SAMPLE_DOCX)
    assert "file_name" in result and "file_type" in result and "full_text" in result and "pages_or_sections" in result
    assert result.get("file_type") == "docx"
    assert "error" not in result, result.get("error")
    _print_result("sample.docx", result)


def test_read_sample_pdf():
    """读取 sample.pdf：不存在则创建后读取，并断言统一结构。"""
    if not os.path.isfile(SAMPLE_PDF):
        _create_sample_pdf(SAMPLE_PDF)
    result = read_document(SAMPLE_PDF)
    assert "file_name" in result and "file_type" in result and "full_text" in result and "pages_or_sections" in result
    assert result.get("file_type") == "pdf"
    assert "error" not in result, result.get("error")
    _print_result("sample.pdf", result)


def test_read_document_dispatch():
    """read_document 根据扩展名正确分发到 PDF/DOCX。"""
    r = read_document("a.PDF")
    assert r.get("file_type") == "pdf"
    r = read_document("b.DOCX")
    assert r.get("file_type") == "docx"
    r = read_document("c.txt")
    assert "error" in r and "Unsupported" in r.get("error", "")
    print("read_document 扩展名分发测试通过。")


if __name__ == "__main__":
    print("Phase4-1 document_reader 简单测试\n")
    test_read_document_dispatch()
    test_read_sample_docx()
    test_read_sample_pdf()
    print("Phase4-1 测试完成。")
